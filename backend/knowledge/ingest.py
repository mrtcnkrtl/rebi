from __future__ import annotations

import json
import os
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, Optional, Sequence

from config import GEMINI_API_KEY, get_logger
from knowledge.db import resolve_postgres_dsn

log = get_logger("knowledge.ingest")

def _sanitize_text_for_pg(text: str) -> str:
    """
    Postgres TEXT cannot contain NUL (0x00) bytes.
    PDFs (and some HTML extractions) may contain embedded NULs.
    """
    if not text:
        return ""
    return str(text).replace("\x00", "").replace("\u0000", "")


def _strip_html(text: str) -> str:
    # very lightweight HTML strip (no external deps)
    text = re.sub(r"(?is)<script.*?>.*?</script>", " ", text)
    text = re.sub(r"(?is)<style.*?>.*?</style>", " ", text)
    text = re.sub(r"(?is)<[^>]+>", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    return _sanitize_text_for_pg(text)


def _read_text_file(path: Path) -> tuple[str, str]:
    ext = path.suffix.lower()
    if ext in {".txt", ".md"}:
        return path.read_text(encoding="utf-8", errors="ignore"), ext.lstrip(".")
    if ext in {".html", ".htm"}:
        raw = path.read_text(encoding="utf-8", errors="ignore")
        return _strip_html(raw), "html"
    if ext == ".pdf":
        try:
            from PyPDF2 import PdfReader  # type: ignore
        except Exception as e:
            raise RuntimeError("PyPDF2 missing; install it to ingest PDFs") from e
        reader = PdfReader(str(path))
        pages = []
        for p in reader.pages:
            try:
                pages.append(p.extract_text() or "")
            except Exception:
                pages.append("")
        return _sanitize_text_for_pg("\n\n".join(pages).strip()), "pdf"
    if ext == ".docx":
        try:
            from docx import Document  # type: ignore
        except Exception as e:
            raise RuntimeError("python-docx missing; pip install python-docx") from e
        doc = Document(str(path))
        parts: list[str] = []
        for para in doc.paragraphs:
            t = (para.text or "").strip()
            if t:
                parts.append(t)
        for table in doc.tables:
            for row in table.rows:
                cells = [(c.text or "").strip() for c in row.cells]
                line = " | ".join(c for c in cells if c)
                if line:
                    parts.append(line)
        return _sanitize_text_for_pg("\n\n".join(parts).strip()), "docx"
    # fallback
    return _sanitize_text_for_pg(path.read_text(encoding="utf-8", errors="ignore")), "other"


def chunk_text(text: str, max_chars: int = 2500, min_chars: int = 400) -> list[str]:
    """
    Simple chunker: paragraph-based, then merges to max_chars.
    This is intentionally deterministic for debugging.
    """
    t = _sanitize_text_for_pg(text or "").strip()
    t = re.sub(r"\n{3,}", "\n\n", t)
    paras = [p.strip() for p in re.split(r"\n\s*\n", t) if p.strip()]
    chunks: list[str] = []
    buf = ""
    for p in paras:
        if not buf:
            buf = p
            continue
        if len(buf) + 2 + len(p) <= max_chars:
            buf = buf + "\n\n" + p
        else:
            if len(buf) >= min_chars:
                chunks.append(buf)
                buf = p
            else:
                # if buffer too small, force merge even if exceeding max_chars a bit
                buf = buf + "\n\n" + p
    if buf.strip():
        chunks.append(buf.strip())
    # final guard: split very large chunks
    out: list[str] = []
    for c in chunks:
        if len(c) <= max_chars * 2:
            out.append(c)
            continue
        # hard split
        for i in range(0, len(c), max_chars):
            out.append(c[i : i + max_chars].strip())
    return [c for c in out if c]


@dataclass
class IngestDoc:
    path: Path
    title: str
    source_type: str
    text: str


_ALLOWED_EXT = {".pdf", ".txt", ".md", ".html", ".htm", ".docx"}


def discover_files(root: Path) -> list[Path]:
    paths = []
    for p in root.rglob("*"):
        if p.is_file() and p.suffix.lower() in _ALLOWED_EXT:
            paths.append(p)
    return sorted(paths)


def _normalize_ingest_paths(paths: Iterable[Path]) -> list[Path]:
    """Mutlak yol, var olan dosyalar; sırayı koru, tekrarları ele."""
    seen: set[str] = set()
    out: list[Path] = []
    for raw in paths:
        p = raw.expanduser().resolve()
        if not p.is_file():
            raise FileNotFoundError(f"Not a file: {raw}")
        if p.suffix.lower() not in _ALLOWED_EXT:
            raise ValueError(f"Unsupported extension (use {_ALLOWED_EXT}): {p}")
        key = str(p)
        if key in seen:
            continue
        seen.add(key)
        out.append(p)
    return out


def _ensure_gemini_client():
    if not GEMINI_API_KEY:
        raise RuntimeError("GEMINI_API_KEY is missing")
    from google import genai  # type: ignore

    return genai.Client(api_key=GEMINI_API_KEY)


def embed_texts_google(
    texts: list[str],
    model: str = "gemini-embedding-001",
    output_dimensionality: int = 768,
) -> list[list[float]]:
    """
    Returns a list of embedding vectors (float lists).
    Uses google-genai client.
    """
    client = _ensure_gemini_client()
    # google-genai supports embed_content; API responses vary slightly by version
    try:
        from google.genai import types  # type: ignore

        resp = client.models.embed_content(
            model=model,
            contents=texts,
            config=types.EmbedContentConfig(output_dimensionality=int(output_dimensionality)),
        )
    except Exception:
        # fallback for older client versions without config typing
        resp = client.models.embed_content(
            model=model,
            contents=texts,
        )
    # Normalize response into list[list[float]] across google-genai versions.
    def _values_from_item(item):
        if item is None:
            return None
        # object forms
        for attr in ("values",):
            v = getattr(item, attr, None)
            if v is not None:
                return v
        emb = getattr(item, "embedding", None)
        if emb is not None:
            v = getattr(emb, "values", None)
            if v is not None:
                return v
            # sometimes embedding itself is list
            if isinstance(emb, (list, tuple)):
                return emb
        # dict forms
        if isinstance(item, dict):
            if "values" in item and item["values"] is not None:
                return item["values"]
            if "embedding" in item and item["embedding"] is not None:
                e = item["embedding"]
                if isinstance(e, dict) and e.get("values") is not None:
                    return e.get("values")
                if isinstance(e, (list, tuple)):
                    return e
        return None

    # google-genai typically returns resp.embeddings (list)
    data = getattr(resp, "embeddings", None)
    if data is None and isinstance(resp, dict):
        data = resp.get("embeddings") or resp.get("data")
    if data is None:
        data = getattr(resp, "data", None) or resp

    # ensure list
    if not isinstance(data, list):
        data = [data]

    embs: list[list[float]] = []
    for item in data:
        v = _values_from_item(item)
        if v is None:
            raise RuntimeError("Unexpected embedding response shape")
        embs.append([float(x) for x in list(v)])
    return embs


def _pg_vector_literal(vec: list[float]) -> str:
    # pgvector accepts: '[1,2,3]'
    return "[" + ",".join(f"{x:.8f}" for x in vec) + "]"


def ingest_paths(
    *,
    user_id: str,
    folder_slug: str,
    folder_title: str,
    paths: Sequence[Path],
    store_raw_text: bool = False,
    embed_model: str = "gemini-embedding-001",
    batch_size: int = 16,
    skip_existing: bool = True,
    replace_existing: bool = False,
) -> dict:
    """
    Yalnızca verilen dosya yollarını işler (tek PDF veya liste).
    skip_existing: aynı mutlak yol (source_url) bu klasörde varsa atla.
    replace_existing: aynı yol varsa önce dokümanı sil (chunk'lar cascade), yeniden yükle.
    """
    dsn = resolve_postgres_dsn()
    if not dsn:
        raise RuntimeError(
            "Postgres DSN missing: set SUPABASE_DATABASE_URL or DATABASE_URL in backend/.env; "
            "or SUPABASE_URL + SUPABASE_DB_PASSWORD (see knowledge/db.py)."
        )
    try:
        import psycopg  # type: ignore
    except Exception as e:
        raise RuntimeError("psycopg is required; install psycopg[binary]") from e

    files = _normalize_ingest_paths(paths)
    log.info("Ingesting %d file(s)", len(files))

    docs: list[IngestDoc] = []
    for p in files:
        text, st = _read_text_file(p)
        text = _sanitize_text_for_pg(text or "").strip()
        if len(text) < 200:
            log.warning("Skipping (too short): %s", p)
            continue
        docs.append(IngestDoc(path=p, title=p.stem, source_type=st, text=text))

    inserted_docs = 0
    skipped_existing = 0
    inserted_document_ids: list[str] = []
    inserted_chunks = 0
    embedded_chunks = 0
    failed_chunks = 0

    with psycopg.connect(dsn, autocommit=True) as conn:
        with conn.cursor() as cur:
            cur.execute(
                """
                insert into public.knowledge_folders (user_id, slug, title)
                values (%s, %s, %s)
                on conflict (user_id, slug) do update set title = excluded.title
                returning id
                """,
                (user_id, folder_slug, folder_title),
            )
            folder_id = cur.fetchone()[0]

            for d in docs:
                source_url = str(d.path.resolve())

                cur.execute(
                    """
                    select id from public.knowledge_documents
                    where user_id = %s and folder_id = %s and source_url = %s
                    limit 1
                    """,
                    (user_id, folder_id, source_url),
                )
                row = cur.fetchone()
                if row:
                    if replace_existing:
                        cur.execute(
                            "delete from public.knowledge_documents where id = %s",
                            (row[0],),
                        )
                    elif skip_existing:
                        log.info("Skip existing (same path): %s", source_url)
                        skipped_existing += 1
                        continue

                cur.execute(
                    """
                    insert into public.knowledge_documents
                      (user_id, folder_id, source_type, title, source_url, raw_text)
                    values (%s, %s, %s, %s, %s, %s)
                    returning id
                    """,
                    (
                        user_id,
                        folder_id,
                        d.source_type,
                        d.title,
                        source_url,
                        d.text if store_raw_text else None,
                    ),
                )
                doc_id = cur.fetchone()[0]
                inserted_docs += 1
                inserted_document_ids.append(str(doc_id))

                chunks = chunk_text(d.text)
                chunks = [_sanitize_text_for_pg(c) for c in chunks]
                for idx, ch in enumerate(chunks):
                    cur.execute(
                        """
                        insert into public.knowledge_chunks
                          (user_id, folder_id, document_id, chunk_index, chunk_text, embed_model, embed_ok)
                        values (%s, %s, %s, %s, %s, %s, false)
                        """,
                        (user_id, folder_id, doc_id, idx, ch, embed_model),
                    )
                inserted_chunks += len(chunks)

                for i in range(0, len(chunks), batch_size):
                    batch = chunks[i : i + batch_size]
                    try:
                        vecs = embed_texts_google(batch, model=embed_model, output_dimensionality=768)
                        if len(vecs) != len(batch):
                            raise RuntimeError("Embedding count mismatch")
                        for j, v in enumerate(vecs):
                            lit = _pg_vector_literal(v)
                            cur.execute(
                                """
                                update public.knowledge_chunks
                                set embedding = %s::vector, embed_ok = true, embed_error = null
                                where document_id = %s and chunk_index = %s
                                """,
                                (lit, doc_id, i + j),
                            )
                        embedded_chunks += len(batch)
                    except Exception as e:
                        failed_chunks += len(batch)
                        err = str(e)[:500]
                        for j in range(len(batch)):
                            cur.execute(
                                """
                                update public.knowledge_chunks
                                set embed_ok = false, embed_error = %s
                                where document_id = %s and chunk_index = %s
                                """,
                                (err, doc_id, i + j),
                            )
                        log.error("Embedding failed for %s batch %d..%d: %s", d.path.name, i, i + len(batch) - 1, err)

    return {
        "folder_slug": folder_slug,
        "folder_title": folder_title,
        "input_paths": [str(d.path.resolve()) for d in docs],
        "files_seen": len(files),
        "inserted_docs": inserted_docs,
        "inserted_document_ids": inserted_document_ids,
        "skipped_existing": skipped_existing,
        "inserted_chunks": inserted_chunks,
        "embedded_chunks": embedded_chunks,
        "failed_chunks": failed_chunks,
    }


def ingest_directory(
    *,
    user_id: str,
    folder_slug: str,
    folder_title: str,
    root_dir: Path,
    store_raw_text: bool = False,
    embed_model: str = "gemini-embedding-001",
    batch_size: int = 16,
    skip_existing: bool = True,
    replace_existing: bool = False,
) -> dict:
    """Klasördeki tüm desteklenen dosyalar; mevcut source_url ile çakışanlar varsayılan olarak atlanır."""
    root_dir = root_dir.expanduser().resolve()
    if not root_dir.is_dir():
        raise NotADirectoryError(f"Not a directory: {root_dir}")
    files = discover_files(root_dir)
    log.info("Discovered %d files under %s", len(files), root_dir)
    result = ingest_paths(
        user_id=user_id,
        folder_slug=folder_slug,
        folder_title=folder_title,
        paths=files,
        store_raw_text=store_raw_text,
        embed_model=embed_model,
        batch_size=batch_size,
        skip_existing=skip_existing,
        replace_existing=replace_existing,
    )
    result["root_dir"] = str(root_dir)
    result["discovered_files"] = len(files)
    return result


if __name__ == "__main__":
    """
    Examples:
      python -m knowledge.ingest --user <uuid> --folder data-pdfs --title "Katalog" --dir ./knowledge-data/data-pdfs
      python -m knowledge.ingest --user <uuid> --folder data-pdfs --title "Katalog" --file ./yeni.pdf
    """
    import argparse

    ap = argparse.ArgumentParser(
        description="Ingest PDF/docx/txt/md/html into knowledge store. Use --file for single files, or --dir for a folder.",
    )
    ap.add_argument("--user", required=True, help="supabase auth user uuid")
    ap.add_argument("--folder", required=True, help="folder slug, e.g. data-pdfs")
    ap.add_argument("--title", required=True, help="folder title")
    src = ap.add_mutually_exclusive_group(required=True)
    src.add_argument("--dir", help="directory to scan recursively")
    src.add_argument(
        "--file",
        action="append",
        dest="files",
        metavar="PATH",
        help="single file to ingest (repeat for multiple); only these paths are processed",
    )
    ap.add_argument("--store-raw", action="store_true", help="store raw_text in documents table")
    ap.add_argument(
        "--replace-existing",
        action="store_true",
        help="if same absolute path already ingested in this folder, delete old doc+chunks and re-ingest",
    )
    ap.add_argument(
        "--no-skip-existing",
        action="store_true",
        help="allow duplicate source_url (not recommended); default is skip same path",
    )
    ap.add_argument("--model", default="gemini-embedding-001", help="Google embedding model id")
    args = ap.parse_args()

    skip_existing = not bool(args.no_skip_existing)
    replace_existing = bool(args.replace_existing)

    if args.dir:
        result = ingest_directory(
            user_id=args.user,
            folder_slug=args.folder,
            folder_title=args.title,
            root_dir=Path(args.dir),
            store_raw_text=bool(args.store_raw),
            embed_model=str(args.model),
            skip_existing=skip_existing,
            replace_existing=replace_existing,
        )
    else:
        result = ingest_paths(
            user_id=args.user,
            folder_slug=args.folder,
            folder_title=args.title,
            paths=[Path(p) for p in (args.files or [])],
            store_raw_text=bool(args.store_raw),
            embed_model=str(args.model),
            skip_existing=skip_existing,
            replace_existing=replace_existing,
        )
    print(json.dumps(result, indent=2, ensure_ascii=False))

